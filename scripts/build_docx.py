"""
Script to generate SentinelDesk_Architecture_and_Agent_Workflow.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def create_document():
    doc = Document()

    # Set page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base styling palette
    NAVY = RGBColor(15, 32, 67)
    BLUE = RGBColor(30, 88, 153)
    DARK_GRAY = RGBColor(50, 50, 50)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = p_title.add_run("SentinelDesk 🛡️ Architecture & Agent Workflow Guide")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY

    # Subtitle
    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("10/10 Enterprise Multi-Agent AI Customer Support Operations Platform")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = BLUE

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- Section 1: Executive Overview ---
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Executive Overview & System Architecture")
    r_h1.font.name = "Arial"
    r_h1.font.color.rgb = NAVY

    p1 = doc.add_paragraph()
    p1.add_run(
        "SentinelDesk is an autonomous multi-agent platform designed to execute end-to-end customer support ticket triage. "
        "Built on Python 3.11+, LangGraph v0.2+, and FastAPI, SentinelDesk converts unstructured user queries into structured, "
        "verifiable, and grounded resolution workflows."
    )

    # --- Section 2: 8-Node State Machine ---
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. The 8-Node Agent State Machine Flow")
    r_h2.font.name = "Arial"
    r_h2.font.color.rgb = NAVY

    # Table for 8 Nodes
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    hdr_titles = ["Step #", "Agent Node Name", "Core Responsibilities"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F2043")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    nodes_data = [
        ("1", "Intake & Security Node", "PII redaction (email, cards, SSN), OWASP LLM01 prompt injection interception, multi-modal OCR attachment ingestion."),
        ("2", "Intent Classifier Node", "Categorizes query into Billing, TechBug, FeatureRequest, AccountAccess, GeneralQuery, or AbusePolicy."),
        ("3", "Urgency & SLA Engine", "Computes urgency score (HOT 0.95, WARM 0.70, COLD 0.25) based on customer tier and keywords."),
        ("4", "Duplicate Search Node", "Queries ChromaDB open ticket vector index to detect duplicate submissions (> 85% similarity)."),
        ("5", "Autonomous ReAct Tool Loop", "Dynamically executes tool chain: lookup_customer_account, verify_transaction, issue_refund, search_knowledge_base."),
        ("6", "RAG Retrieval Node", "Local bge-small-en-v1.5 embeddings → ChromaDB vector query across faq, manual, policy, and guide collections."),
        ("7", "Draft Resolution Node", "Generates customer-facing resolution text grounded strictly in retrieved KB chunks with source citations."),
        ("8", "Decision & Gate Node", "Confidence Gate check (>= 0.75 SOLVED; < 0.75 or policy rules ESCALATED to Tier 2 / Compliance)."),
    ]

    for step_num, node_name, desc in nodes_data:
        row_cells = table.add_row().cells
        row_cells[0].text = step_num
        row_cells[1].text = node_name
        row_cells[2].text = desc

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- Section 3: Autonomous ReAct Engine ---
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. Autonomous ReAct Tool Execution & Self-Correction")
    r_h3.font.name = "Arial"
    r_h3.font.color.rgb = NAVY

    doc.add_paragraph(
        "SentinelDesk employs an autonomous ReAct loop (Thought → Action → Observation → Reflexion). "
        "The agent dynamically selects tools from its executable registry:"
    )

    tools_list = [
        "lookup_customer_account(email): Retrieves customer account history, subscription tier, and lifetime value.",
        "verify_transaction(transaction_id): Queries ledger to verify double charges and refund eligibility.",
        "issue_refund(customer_id, amount, reason): Autonomous refund execution for verified billing disputes.",
        "search_knowledge_base(query, top_k): Queries ChromaDB collections for technical documentation.",
    ]
    for tool_item in tools_list:
        doc.add_paragraph(tool_item, style="List Bullet")

    # --- Section 4: Security & Resilience ---
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. Security Guardrails & Multi-Model Failover")
    r_h4.font.name = "Arial"
    r_h4.font.color.rgb = NAVY

    doc.add_paragraph(
        "• OWASP LLM Top 10 Defense: Strict input regex sanitization & injection payload isolation.\n"
        "• Multi-Model Failover Pool: Automatic fallback (Gemini 2.0 Flash → Gemini 2.0 Flash Lite → Groq → OpenRouter) on 429 Rate Limits.\n"
        "• Real-Time WebSockets: Live telemetry streamed over ws://localhost:8000/ws/live-triage."
    )

    # Save document
    output_path = "SentinelDesk_Architecture_and_Agent_Workflow.docx"
    doc.save(output_path)
    print(f"Successfully generated: {os.path.abspath(output_path)}")


if __name__ == "__main__":
    create_document()
